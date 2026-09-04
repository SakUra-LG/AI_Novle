from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs_pop_king_v6_compiled_story_first_500"
CHAPTER_DIR = OUT / "chapters"
MD_PATH = OUT / "麦珂重生小说_中期成果_第001-210章.md"
DOCX_PATH = OUT / "麦珂重生小说_中期成果_第001-210章.docx"


def set_run_font(run, name, size, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_or_paragraph_font(paragraph, name, size, bold=False, color=None):
    for run in paragraph.runs:
        set_run_font(run, name, size, bold, color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, "Calibri", 9, color="777777")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clean_body(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return [line.strip() for line in text.split("\n") if line.strip()]


def main():
    chapters = []
    missing = []
    for chapter in range(1, 211):
        path = CHAPTER_DIR / f"chapter_{chapter:03d}.txt"
        if not path.is_file():
            missing.append(chapter)
            continue
        chapters.append((chapter, clean_body(path.read_text(encoding="utf-8"))))
    if missing:
        raise RuntimeError(f"missing chapters: {missing}")

    md_parts = [
        "# 麦珂重生小说（中期成果）",
        "",
        "正文第1—210章整合稿",
        "",
        "---",
        "",
    ]
    for chapter, paragraphs in chapters:
        md_parts.append(f"## 第{chapter}章")
        md_parts.append("")
        md_parts.extend(paragraphs)
        md_parts.extend(["", "---", ""])
    MD_PATH.write_text("\n".join(md_parts), encoding="utf-8", newline="\n")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h1 = doc.styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor.from_string("2E74B5")
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("麦珂重生小说")
    set_run_font(run, "黑体", 24, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("中期成果整合稿｜第1—210章")
    set_run_font(run, "宋体", 12, color="555555")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(18)
    run = note.add_run("正文按现有正式章节文件合并，章节号已统一标注。")
    set_run_font(run, "宋体", 10, color="777777")
    doc.add_page_break()

    for index, (chapter, paragraphs) in enumerate(chapters):
        heading = doc.add_paragraph(style="Heading 1")
        heading.add_run(f"第{chapter}章")
        for paragraph_text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.keep_together = False
            run = p.add_run(paragraph_text)
            set_run_font(run, "宋体", 11)
        if index < len(chapters) - 1:
            doc.add_page_break()

    doc.core_properties.title = "麦珂重生小说——中期成果整合稿（第1—210章）"
    doc.core_properties.subject = "小说正文中期成果"
    doc.core_properties.author = ""
    doc.save(DOCX_PATH)
    print(f"written markdown: {MD_PATH}")
    print(f"written docx: {DOCX_PATH}")
    print(f"chapters: {len(chapters)}")


if __name__ == "__main__":
    main()
