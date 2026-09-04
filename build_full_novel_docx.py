from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = ROOT / "bert_excitation_train"
CHAPTER_DIR = PROJECT_ROOT / "outputs_pop_king_v6_compiled_story_first_500" / "chapters"
OUT = PROJECT_ROOT / "outputs_pop_king_v6_compiled_story_first_500" / "全书500章正文.docx"


def set_font(run, name="宋体", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="宋体", size=11, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor(*color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    return run


def add_page_break(paragraph):
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def chapter_number(path):
    return int(re.search(r"chapter_(\d+)", path.stem).group(1))


def main():
    files = sorted(CHAPTER_DIR.glob("chapter_*.txt"), key=chapter_number)
    expected = list(range(1, 501))
    actual = [chapter_number(p) for p in files]
    if actual != expected:
        missing = sorted(set(expected) - set(actual))
        extra = sorted(set(actual) - set(expected))
        raise SystemExit(f"chapter set invalid: missing={missing}, extra={extra}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, "宋体", 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "黑体", 16, (46, 116, 181))
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    for style_name in ("Heading 2", "Heading 3"):
        st = doc.styles[style_name]
        set_style_font(st, "黑体", 13 if style_name == "Heading 2" else 12, (46, 116, 181) if style_name == "Heading 2" else (31, 77, 120))

    # Quiet running furniture for a long-form reading document.
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("500章正文合集")
    set_font(hr, "黑体", 9, (110, 110, 110))
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("第 ")
    set_font(fr, "宋体", 9, (110, 110, 110))
    add_page_field(fp)
    fr2 = fp.add_run(" 页")
    set_font(fr2, "宋体", 9, (110, 110, 110))

    # Editorial-cover opening, with a restrained, novel-appropriate override.
    for _ in range(7):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("长篇正文")
    set_font(r, "黑体", 12, (122, 90, 0), bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("500章正文合集")
    set_font(r, "黑体", 30, (32, 55, 72), bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("第一章—第五百章")
    set_font(r, "宋体", 15, (43, 81, 99))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("整合版正文")
    set_font(r, "宋体", 10.5, (100, 100, 100))
    add_page_break(doc.add_paragraph())

    for idx, path in enumerate(files, start=1):
        chapter_text = path.read_text(encoding="utf-8").replace("\r\n", "\n").replace("\r", "\n")
        heading = doc.add_paragraph(style="Heading 1")
        heading.add_run(f"第{idx}章")

        blocks = re.split(r"\n\s*\n", chapter_text.strip())
        for block in blocks:
            lines = [line.strip() for line in block.split("\n") if line.strip()]
            if not lines:
                continue
            text = "".join(lines)
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.first_line_indent = Inches(0.3)
            para.add_run(text)

        if idx != 500:
            add_page_break(doc.add_paragraph())

    doc.core_properties.title = "500章正文合集"
    doc.core_properties.subject = "第1—500章整合正文"
    doc.core_properties.author = ""
    doc.core_properties.comments = "由正式章节文件按编号顺序整合生成"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"created {OUT}")
    print(f"chapters={len(files)}")


if __name__ == "__main__":
    main()
