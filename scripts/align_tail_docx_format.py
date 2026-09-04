import argparse
import copy
import os
import re
import tempfile
from pathlib import Path

from docx import Document


CHAPTER_RE = re.compile(r"^\s*第\s*(\d+)\s*章(?:\s|[:：]|$)")


def replace_ppr(target, source):
    target_p = target._p
    old = target_p.pPr
    if old is not None:
        target_p.remove(old)
    if source._p.pPr is not None:
        target_p.insert(0, copy.deepcopy(source._p.pPr))


def clear_run_formatting(paragraph):
    for run in paragraph.runs:
        rpr = run._r.rPr
        if rpr is not None:
            run._r.remove(rpr)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--heading-start", type=int, default=299)
    parser.add_argument("--body-start", type=int, default=386)
    parser.add_argument("--end", type=int, default=500)
    args = parser.parse_args()

    doc = Document(args.docx)
    headings = {}
    first_body = {}
    current = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = CHAPTER_RE.match(text)
        if match:
            current = int(match.group(1))
            headings[current] = paragraph
            continue
        if current is not None and text and current not in first_body:
            first_body[current] = paragraph

    if 298 not in headings or 298 not in first_body:
        raise RuntimeError("chapter 298 format reference not found")
    missing = [n for n in range(args.heading_start, args.end + 1) if n not in headings]
    if missing:
        raise RuntimeError(f"missing chapter headings: {missing}")

    heading_ref = headings[298]
    body_ref = first_body[298]
    current = None
    heading_count = 0
    body_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = CHAPTER_RE.match(text)
        if match:
            current = int(match.group(1))
            if args.heading_start <= current <= args.end:
                replace_ppr(paragraph, heading_ref)
                paragraph.style = heading_ref.style
                clear_run_formatting(paragraph)
                heading_count += 1
            continue
        if current is not None and args.body_start <= current <= args.end and text:
            replace_ppr(paragraph, body_ref)
            paragraph.style = body_ref.style
            clear_run_formatting(paragraph)
            body_count += 1

    temp_fd, temp_name = tempfile.mkstemp(prefix=args.docx.stem + "_", suffix=".docx", dir=args.docx.parent)
    os.close(temp_fd)
    temp_path = Path(temp_name)
    try:
        doc.save(temp_path)
        # Reopen before replacing the authoritative file to prove the package is valid.
        Document(temp_path)
        os.replace(temp_path, args.docx)
    finally:
        if temp_path.exists():
            temp_path.unlink()
    print({"headings_aligned": heading_count, "body_paragraphs_aligned": body_count, "path": str(args.docx)})


if __name__ == "__main__":
    main()
