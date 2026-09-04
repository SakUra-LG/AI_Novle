from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


CHAPTER_RE = re.compile(r"^\s*第\s*([0-9一二三四五六七八九十百零〇两]+)\s*章(?:\s|[:：]|$)")


def chinese_number(text: str) -> int | None:
    if text.isdigit():
        return int(text)
    digits = {"零": 0, "〇": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
              "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    units = {"十": 10, "百": 100}
    total = 0
    current = 0
    for char in text:
        if char in digits:
            current = digits[char]
        elif char in units:
            total += (current or 1) * units[char]
            current = 0
        else:
            return None
    return total + current


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--chapters", default="1-20")
    parser.add_argument("--out", type=Path)
    parser.add_argument("--compact", action="store_true")
    parser.add_argument("--summary", action="store_true")
    args = parser.parse_args()

    lo, hi = (int(v) for v in args.chapters.split("-", 1))
    doc = Document(args.docx)
    chapters: dict[int, dict] = {}
    current: int | None = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        match = CHAPTER_RE.match(text)
        if match:
            number = chinese_number(match.group(1))
            if number is not None:
                current = number
                chapters[current] = {
                    "number": current,
                    "heading": text,
                    "heading_index": index,
                    "heading_style": paragraph.style.name if paragraph.style else None,
                    "paragraphs": [],
                }
                continue
        if current is not None:
            chapters[current]["paragraphs"].append({
                "index": index,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else None,
            })

    selected = []
    for number in range(lo, hi + 1):
        if number not in chapters:
            continue
        item = chapters[number]
        body = "\n".join(p["text"] for p in item["paragraphs"] if p["text"].strip())
        item["char_count"] = len(re.sub(r"\s+", "", body))
        if args.summary:
            item = {
                "number": number,
                "heading": item["heading"],
                "heading_index": item["heading_index"],
                "char_count": item["char_count"],
                "paragraph_count": len([p for p in item["paragraphs"] if p["text"].strip()]),
                "opening": next((p["text"] for p in item["paragraphs"] if p["text"].strip()), "")[:80],
                "ending": next((p["text"] for p in reversed(item["paragraphs"]) if p["text"].strip()), "")[-100:],
            }
        selected.append(item)

    payload = {"selected": selected}
    if not args.compact:
        payload = {
            "paragraph_count": len(doc.paragraphs),
            "chapter_count": len(chapters),
            "chapter_numbers": sorted(chapters),
            "selected": selected,
        }
    content = json.dumps(payload, ensure_ascii=False, indent=2)
    if args.out:
        args.out.write_text(content, encoding="utf-8")
    else:
        print(content)


if __name__ == "__main__":
    main()
