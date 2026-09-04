from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document


DOCX = Path(
    r"D:\Study\College\Scientific research\张颖——AI小说自动生成\张颖——AI小说自动生成"
    r"\bert_excitation_train\AI_Novle\第一版本\全书500章正文_386-500人工修订终稿.docx"
)

HEADING_RE = re.compile(r"^\s*第\s*(\d+)\s*章(?:\s|[:：]|$)")


def set_text_preserving_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(DOCX)
    updated: list[int] = []
    headings: list[int] = []

    for paragraph in doc.paragraphs:
        current = paragraph.text.strip()
        match = HEADING_RE.match(current)
        if not match:
            continue
        chapter = int(match.group(1))
        headings.append(chapter)
        expected = f"第{chapter}章"
        if current != expected:
            set_text_preserving_first_run(paragraph, expected)
            updated.append(chapter)

    if headings != list(range(1, 501)):
        raise RuntimeError(f"Chapter sequence is invalid: {headings[:5]} ... {headings[-5:]}")
    if updated != list(range(299, 501)):
        raise RuntimeError(f"Unexpected heading edit scope: {updated[:5]} ... {updated[-5:]}")

    temp = DOCX.with_name(DOCX.stem + "_heading_tmp.docx")
    doc.save(temp)

    check = Document(temp)
    final_headings: list[tuple[int, str, str | None]] = []
    for paragraph in check.paragraphs:
        current = paragraph.text.strip()
        match = HEADING_RE.match(current)
        if match:
            final_headings.append(
                (
                    int(match.group(1)),
                    current,
                    paragraph.style.name if paragraph.style else None,
                )
            )

    if len(final_headings) != 500:
        raise RuntimeError(f"Expected 500 headings, found {len(final_headings)}")
    for chapter, text, style in final_headings:
        if text != f"第{chapter}章":
            raise RuntimeError(f"Chapter {chapter} still has extra title text: {text!r}")
        if style != "Heading 1":
            raise RuntimeError(f"Chapter {chapter} heading style changed: {style!r}")

    os.replace(temp, DOCX)
    print({"updated": len(updated), "range": [updated[0], updated[-1]], "headings": 500})


if __name__ == "__main__":
    main()
