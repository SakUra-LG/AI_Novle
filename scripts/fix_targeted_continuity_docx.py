from __future__ import annotations

import json
import os
import re
import tempfile
from pathlib import Path

from docx import Document


DOCX = Path(r"D:\Study\College\Scientific research\张颖——AI小说自动生成\张颖——AI小说自动生成\bert_excitation_train\AI_Novle\第一版本\全书500章正文_386-500人工修订终稿.docx")
SYNOPSES = Path(r"D:\Study\College\Scientific research\张颖——AI小说自动生成\张颖——AI小说自动生成\bert_excitation_train\AI_Novle\bert_excitation_train\outputs_pop_king_v6_compiled_story_first_500\chapter_synopses_v5_qwen_500.json")
CLUSTERS = Path(r"D:\Study\College\Scientific research\张颖——AI小说自动生成\张颖——AI小说自动生成\bert_excitation_train\AI_Novle\bert_excitation_train\outputs_pop_king_v6_compiled_story_first_500\event_clusters_v2.json")

CHAPTER_RE = re.compile(r"^\s*第\s*(\d+)\s*章")


def replacements_for(chapter: int) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    if 29 <= chapter <= 50:
        pairs.append(("十一岁", "十三岁"))
        if chapter == 45:
            pairs.append(("十二岁的自己", "十三岁的自己"))
    if 53 <= chapter <= 54:
        pairs.append(("十一岁", "十二岁"))
    if 71 <= chapter <= 80:
        pairs.append(("十一岁", "十四岁"))
    if chapter in (79, 99):
        pairs.append(("加州", "银湾州"))
    if chapter == 95:
        pairs.append(("有人拿出手机悄悄拍照", "有人举起便携式相机悄悄拍照"))
    if chapter == 96:
        pairs.append(("苏菲亚点了点头，拿出手机拨通了电话", "苏菲亚点了点头，抓起柜台公用电话拨通了号码"))
    if chapter == 140:
        pairs.append(("有人举起了手机（虽然这时代很少见，但有人带了便携式录音机）", "有人举起了便携式录音机，也有人举起相机"))
    if chapter == 193:
        pairs.append(("他掏出手机，拨通了一个号码", "他抓起桌上电话，拨通了一个号码"))
    if chapter == 288:
        pairs.append(("社交媒体和街头巷尾", "报纸读者来信、广播热线和街头巷尾"))
    if chapter in (367, 368):
        pairs.append(("道歉不买热搜", "道歉不买头版"))
        pairs.append(("热搜庆典", "门户首页庆典"))
    return pairs


def replace_in_paragraph(paragraph, old: str, new: str) -> int:
    count = paragraph.text.count(old)
    if not count:
        return 0
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
    if old in paragraph.text:
        full = paragraph.text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = full
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(full)
    return count


def atomic_write_json(path: Path, data) -> None:
    fd, name = tempfile.mkstemp(prefix=path.stem + "_", suffix=".json", dir=path.parent)
    os.close(fd)
    temp = Path(name)
    try:
        temp.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        json.loads(temp.read_text(encoding="utf-8"))
        os.replace(temp, path)
    finally:
        if temp.exists():
            temp.unlink()


def replace_recursive(value, pairs: list[tuple[str, str]]):
    if isinstance(value, str):
        for old, new in pairs:
            value = value.replace(old, new)
        return value
    if isinstance(value, list):
        return [replace_recursive(item, pairs) for item in value]
    if isinstance(value, dict):
        return {key: replace_recursive(item, pairs) for key, item in value.items()}
    return value


def cluster_span(item: dict) -> tuple[int, int] | None:
    span = item.get("chapter_span")
    if isinstance(span, list) and len(span) == 2:
        return int(span[0]), int(span[1])
    start = item.get("cluster_span_start") or item.get("chapter_start")
    end = item.get("cluster_span_end") or item.get("chapter_end")
    if start is not None and end is not None:
        return int(start), int(end)
    return None


def main() -> None:
    doc = Document(DOCX)
    current = None
    change_log: dict[int, dict[str, int]] = {}
    for paragraph in doc.paragraphs:
        match = CHAPTER_RE.match(paragraph.text.strip())
        if match:
            current = int(match.group(1))
        if current is None:
            continue
        for old, new in replacements_for(current):
            count = replace_in_paragraph(paragraph, old, new)
            if count:
                change_log.setdefault(current, {})[f"{old}->{new}"] = change_log.setdefault(current, {}).get(f"{old}->{new}", 0) + count

    temp_docx = DOCX.with_name(DOCX.stem + "_continuity_tmp.docx")
    doc.save(temp_docx)
    Document(temp_docx)
    os.replace(temp_docx, DOCX)

    synopses = json.loads(SYNOPSES.read_text(encoding="utf-8"))
    for item in synopses:
        chapter = int(item["chapter_id"])
        pairs = replacements_for(chapter)
        if pairs:
            updated = replace_recursive(item, pairs)
            item.clear()
            item.update(updated)
    atomic_write_json(SYNOPSES, synopses)

    clusters_raw = json.loads(CLUSTERS.read_text(encoding="utf-8"))
    cluster_list = clusters_raw if isinstance(clusters_raw, list) else clusters_raw.get("event_clusters", clusters_raw.get("clusters", clusters_raw.get("data", [])))
    for item in cluster_list:
        span = cluster_span(item)
        if not span:
            continue
        pairs: list[tuple[str, str]] = []
        for chapter in range(span[0], span[1] + 1):
            for pair in replacements_for(chapter):
                if pair not in pairs:
                    pairs.append(pair)
        if pairs:
            updated = replace_recursive(item, pairs)
            item.clear()
            item.update(updated)
    atomic_write_json(CLUSTERS, clusters_raw)

    print(json.dumps({"chapters_changed": sorted(change_log), "changes": change_log}, ensure_ascii=False))


if __name__ == "__main__":
    main()
