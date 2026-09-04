import argparse
import copy
import hashlib
import json
import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HEADING_RE = re.compile(r"^\s*第\s*(\d+)\s*章(?:\s|[:：]|$)")


def element_text(node):
    return "".join(part.text or "" for part in node.iter(qn("w:t"))).strip()


def heading_number(node):
    if node.tag != qn("w:p"):
        return None
    match = HEADING_RE.match(element_text(node))
    return int(match.group(1)) if match else None


def make_paragraph(text, template_ppr):
    paragraph = OxmlElement("w:p")
    if template_ppr is not None:
        paragraph.append(copy.deepcopy(template_ppr))
    if text:
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        paragraph.append(run)
    return paragraph


def chapter_nodes(document, start, end_exclusive):
    body = document._body._element
    active = False
    result = []
    for node in body.iterchildren():
        number = heading_number(node)
        if number == start:
            active = True
        if number == end_exclusive:
            break
        if active:
            result.append(node)
    return result


def text_digest(document, start, end):
    pieces = []
    current = None
    for paragraph in document.paragraphs:
        text = paragraph.text
        match = HEADING_RE.match(text.strip())
        if match:
            current = int(match.group(1))
        if current is not None and start <= current <= end:
            pieces.append(text)
    return hashlib.sha256("\n".join(pieces).encode("utf-8")).hexdigest()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("target", type=Path)
    parser.add_argument("recovery_source", type=Path)
    parser.add_argument("chapter42_draft", type=Path)
    args = parser.parse_args()

    target = Document(args.target)
    source = Document(args.recovery_source)
    draft = json.loads(args.chapter42_draft.read_text(encoding="utf-8"))["42"]

    source_42 = chapter_nodes(source, 42, 43)
    source_43_298 = chapter_nodes(source, 43, 299)
    if not source_42 or heading_number(source_42[0]) != 42:
        raise RuntimeError("chapter 42 heading not found in recovery source")
    if not source_43_298 or heading_number(source_43_298[0]) != 43:
        raise RuntimeError("chapter 43 heading not found in recovery source")

    target_body = target._body._element
    target_299 = next((n for n in target_body.iterchildren() if heading_number(n) == 299), None)
    if target_299 is None:
        raise RuntimeError("target chapter 299 heading not found")

    heading_42 = copy.deepcopy(source_42[0])
    template_ppr = None
    for node in source_42[1:]:
        if node.tag == qn("w:p"):
            template_ppr = node.find(qn("w:pPr"))
            break

    insertion = [heading_42]
    insertion.extend(make_paragraph(text, template_ppr) for text in draft)
    insertion.append(make_paragraph("", template_ppr))
    insertion.extend(copy.deepcopy(node) for node in source_43_298)
    for node in insertion:
        target_299.addprevious(node)

    fd, temp_name = tempfile.mkstemp(prefix=args.target.stem + "_", suffix=".docx", dir=args.target.parent)
    os.close(fd)
    temp = Path(temp_name)
    try:
        target.save(temp)
        check = Document(temp)
        numbers = []
        for paragraph in check.paragraphs:
            match = HEADING_RE.match(paragraph.text.strip())
            if match:
                numbers.append(int(match.group(1)))
        if numbers != list(range(1, 501)):
            raise RuntimeError(f"chapter sequence invalid: count={len(numbers)}")
        recovered_digest = text_digest(check, 43, 298)
        source_digest = text_digest(source, 43, 298)
        if recovered_digest != source_digest:
            raise RuntimeError("recovered chapters 43-298 do not match source")
        os.replace(temp, args.target)
    finally:
        if temp.exists():
            temp.unlink()

    print(json.dumps({
        "chapter_count": 500,
        "restored_range": [43, 298],
        "restored_digest_matches": True,
        "chapter42_paragraphs": len(draft),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
